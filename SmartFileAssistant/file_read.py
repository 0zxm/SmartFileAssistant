import os
import re
import shutil
import tempfile
import subprocess
from functools import wraps

import pandas as pd
from docx import Document
from PyPDF2 import PdfReader

DEFAULT_ANTIWORD_PATH = r"E:\VSCodeProject\antiword\antiword.exe"
DEFAULT_MAPPING_PATH = r"E:\VSCodeProject\antiword\UTF-8.txt"
ANTIWORD_CMD = os.environ.get("ANTIWORD_CMD", DEFAULT_ANTIWORD_PATH)
ANTIWORD_MAPPING_FILE = os.environ.get("ANTIWORD_MAPPING_FILE", DEFAULT_MAPPING_PATH)

def validate_file_exists(func):
    @wraps(func)
    def wrapper(self, file_path, *args, **kwargs):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件 {file_path} 不存在")
        return func(self, file_path, *args, **kwargs)
    return wrapper


class FileReader:

    def __init__(self):
        pass
        

    def __getattr__(self, attr_name):
        if attr_name.startswith("__") and attr_name.endswith("__"):
            raise AttributeError(attr_name)
        print("要查找的属性", attr_name)
        raise NotImplementedError(f"没有实现阅读此类文件的方法{attr_name}")
    
    @validate_file_exists
    def read_docx_func(self, file_path:str):
        """增强版 docx 读取（保留表格/页眉页脚）"""
        content_parts = []
        
        try:
            doc = Document(file_path)
            
            # 正文段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text.strip())
            
            # 表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(" | ".join(row_text))
            
            # 页眉页脚
            for section in doc.sections:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text.strip())
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text.strip())
            
            content = "\n".join(content_parts)
            
            if not content.strip():
                print(f"⚠️  警告：{file_path} 未读取到任何内容")
            
            return content
            
        except Exception as e:
            print(f"❌ 读取失败：{file_path}, 错误：{str(e)}")
            return ""

    @validate_file_exists
    def read_pdf_func(self, file_path:str):
        reader = PdfReader(file_path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except:
                print(f"❌ PDF加密无法读取：{file_path}")
                return None
        content = "\n".join([page.extract_text() or "" for page in reader.pages])
        return content
    
    @validate_file_exists
    def read_txt_func(self, file_path:str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        return content
    
    @validate_file_exists
    def read_md_func(self, file_path:str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        md_re_list = ["^#+", ]
        for pattern in md_re_list: # 将内容中的标题语法删除
            re.sub(pattern, "", string=content, flags=re.MULTILINE)

        return content

    @validate_file_exists
    def read_xlsx_func(self, file_path: str) -> str:
        try:
            all_text = []
            xls = pd.ExcelFile(file_path, engine='openpyxl')
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name, header=None, engine='openpyxl')
                # 去除完全空的行
                df = df.dropna(how='all')
                # 拼接有效单元格文本
                sheet_text = ' '.join(
                    df.astype(str)
                    .replace('nan', '', regex=True)
                    .stack()
                    .str.strip()
                    .tolist()
                )
                if sheet_text:
                    all_text.append(f'【工作表{sheet_name}】{sheet_text}')

            return ' '.join(all_text).strip()
        except Exception as e:
            print(f"读取Excel失败：{e}")
            return ""

    @validate_file_exists
    def read_doc_func(self, file_path: str) -> str:
        """使用 antiword 命令将 .doc 转换为纯文本（最终修复版：移除-m参数）"""
        # 1. 检查antiword是否存在（兼容手动指定的绝对路径）
        antiword_path = shutil.which(ANTIWORD_CMD) or ANTIWORD_CMD
        if not os.path.exists(antiword_path):
            raise FileNotFoundError(
                "未检测到 antiword，请安装并确保其在 PATH 中，或设置 ANTIWORD_CMD 环境变量。"
            )

        # 2. 检查映射文件（仅检查默认目录，不手动指定-m）
        # Windows版antiword默认查找路径：C:\Users\Canway\antiword\UTF-8.txt
        user_antiword_dir = os.path.join(os.path.expanduser("~"), "antiword")
        default_mapping = os.path.join(user_antiword_dir, "UTF-8.txt")
        
        if not os.path.exists(default_mapping):
            # 自动创建目录，提示用户放文件
            os.makedirs(user_antiword_dir, exist_ok=True)
            raise FileNotFoundError(
                f"请将UTF-8.txt文件放入 {user_antiword_dir} 目录后重试！\n"
                f"UTF-8.txt内容只需包含：\n"
                "# UTF-8 to Unicode\n# This file is a dummy.\n# The conversion is done algorithmicly, not by a table look-up."
            )

        # 3. 配置环境变量（关键：HOME指向用户目录）
        env = os.environ.copy()
        env["HOME"] = os.path.expanduser("~")  # 强制指向C:\Users\Canway
        env.pop("ANTIWORDHOME", None)  # 移除干扰项

        # 4. 临时目录处理中文/空格路径（核心）
        with tempfile.TemporaryDirectory() as temp_dir:
            # 生成纯英文临时文件名
            temp_file_name = "temp_doc_convert.doc"
            temp_file_path = os.path.join(temp_dir, temp_file_name)
            
            # 复制原文件到临时目录
            try:
                shutil.copy2(file_path, temp_file_path)
            except Exception as e:
                print(f"❌ 复制文件到临时目录失败：{file_path}, 错误：{str(e)}")
                return ""

            # 5. 构建命令（彻底移除-m参数！！！）
            try:
                cmd = [
                    antiword_path,
                    "-t",  # 纯文本输出
                    "-w", "0",  # 取消行宽限制
                    temp_file_name  # 仅传临时文件名（cwd已指定临时目录）
                ]
                
                completed = subprocess.run(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    check=False,
                    env=env,
                    cwd=temp_dir,  # 固定工作目录
                    creationflags=0x08000000,  # 隐藏cmd窗口
                    shell=False  # 禁用shell避免路径解析错误
                )
            except Exception as e:
                print(f"❌ 调用 antiword 失败：{file_path}, 错误：{str(e)}")
                return ""

        # 6. 处理转换结果
        if completed.returncode != 0:
            err_msg = completed.stderr.decode("gbk", errors="ignore").strip() or completed.stderr.decode("utf-8", errors="ignore").strip()
            raise RuntimeError(f"antiword 转换失败：{err_msg or '未知错误'}")

        # 7. 解码输出内容
        stdout_bytes = completed.stdout or b""
        try:
            text = stdout_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            text = stdout_bytes.decode("gbk", errors="ignore").strip()
        
        return text


if __name__ == '__main__':
    fr = FileReader()
    # print(fr.x)

    content = fr.read_xlsx_func(r'C:\Users\Canway\Downloads\信通院算力需求开发排期【内部】.xlsx')
    print(content)